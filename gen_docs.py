from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "4F46E5")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if r_idx % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EEF2FF")
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

# ── COVER ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MediCluster")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(79, 70, 229)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Patient Health Risk Segregation Platform")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dp.add_run("Technical Documentation  |  " + datetime.date.today().strftime("%B %Y")).font.size = Pt(11)
doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
heading("1. Executive Summary", 1, (79, 70, 229))
para(
    "MediCluster is a full-stack AI-powered clinical decision support platform that clusters patients "
    "by health risk, analyzes medical images, processes clinical notes with NLP, forecasts vital signs, "
    "dispatches ambulances intelligently, and enables real-time AI-assisted triage."
)
para("Key highlights:", bold=True)
bullet("4 unsupervised ML clustering algorithms: K-Means, DBSCAN, Hierarchical, GMM")
bullet("LungLens: Claude Vision imaging model for any medical image type")
bullet("CliniQ: multilingual AI assistant (English, Hindi, Kannada, Telugu, Tamil)")
bullet("ARIA: AI-powered ambulance dispatch with live maps and driver alerting")
bullet("Clinical NLP: NER, ICD-10 coding, drug interaction checking, vital forecasting")
bullet("React frontend + Node.js API gateway + Python Flask ML engine + MongoDB")
doc.add_page_break()

# ── 2. ARCHITECTURE ───────────────────────────────────────────────────────────
heading("2. System Architecture", 1, (79, 70, 229))
para("MediCluster is a three-tier application with four services:")
add_table(
    ["Layer", "Technology", "Port", "Responsibility"],
    [
        ["Frontend",    "React 18 + Vite + Tailwind CSS",      "3000", "All user-facing pages and dashboards"],
        ["API Gateway", "Node.js 25 + Express 4 + Mongoose",   "5000", "REST API, Claude API calls, MongoDB persistence"],
        ["ML Engine",   "Python 3.11 + Flask 3 + scikit-learn","8000", "Clustering, NLP, imaging, forecasting, anomaly detection"],
        ["Database",    "MongoDB 7 (Windows service)",          "27017","Patient data, cluster results, media (GridFS), reminders"],
    ],
    [1.3, 2.3, 0.7, 2.3]
)
heading("2.1 Request Flow", 2)
bullet("Browser sends REST requests to Node.js backend (port 5000)")
bullet("Backend validates, persists metadata to MongoDB, then proxies ML requests to Flask engine (port 8000)")
bullet("Claude API calls are made from the Node.js backend via Anthropic SDK (shared client, 120s timeout)")
bullet("Media files stored in MongoDB GridFS; base64 copies forwarded to ML engine for analysis")
doc.add_page_break()

# ── 3. FEATURES ───────────────────────────────────────────────────────────────
heading("3. Feature Modules", 1, (79, 70, 229))

heading("3.1 Patient Risk Clustering", 2)
para("Upload a CSV of patient vitals; the platform clusters patients into risk tiers (Critical, High, Moderate, Low).")
add_table(
    ["Algorithm", "Best For", "Key Parameter"],
    [
        ["K-Means",      "Well-separated spherical clusters",         "k (number of clusters)"],
        ["DBSCAN",       "Arbitrary shape clusters, noise isolation",  "eps, min_samples"],
        ["Hierarchical", "Hierarchical relationships + dendrogram",    "linkage (ward/complete/average)"],
        ["GMM",          "Soft/probabilistic cluster membership",      "n_components, covariance type"],
    ],
    [1.5, 2.8, 2.3]
)
bullet("AutoML: sweeps K 2-10, picks best via Elbow + Silhouette")
bullet("UMAP and t-SNE dimensionality reduction for 2D visualization")
bullet("SHAP explainability: per-patient feature contributions")
bullet("Isolation Forest anomaly detection flags statistical outliers")

heading("3.2 Medical Imaging (LungLens + DenseNet)", 2)
para("Supports chest X-ray and DICOM analysis using five DenseNet/ResNet models plus LungLens.")
add_table(
    ["Model", "Label", "Dataset"],
    [
        ["densenet121-res224-chex",  "CheXNet (recommended)", "Stanford CheXpert 224k"],
        ["densenet121-res224-all",   "DenseNet121 All",       "CheXpert + NIH + MIMIC + PadChest"],
        ["resnet50-res512-all",      "ResNet50 High-res",     "Multi-dataset, 512px"],
        ["densenet121-res224-nih",   "DenseNet121 NIH",       "NIH ChestX-ray14"],
        ["densenet121-res224-pc",    "DenseNet121 PadChest",  "PadChest"],
        ["claude-vision (LungLens)", "LungLens",              "Anthropic Claude Vision (any image)"],
    ],
    [2.2, 1.8, 2.6]
)
bullet("Each finding: label, confidence %, severity (HIGH/MODERATE/LOW), cause, medications, prevention")
bullet("Grad-CAM heatmaps show which image region triggered each finding")
bullet("AI Deep Explanation: Claude generates plain-English paragraph explanations")

heading("3.3 CliniQ - Multilingual AI Assistant", 2)
para("Vision-capable AI chat. Upload any medical image and get analysis in 5 languages.")
add_table(
    ["Language", "Script",      "Code"],
    [
        ["English", "Latin",      "en"],
        ["Hindi",   "Devanagari", "hi"],
        ["Kannada", "Kannada",    "kn"],
        ["Telugu",  "Telugu",     "te"],
        ["Tamil",   "Tamil",      "ta"],
    ],
    [1.5, 1.5, 1.0]
)

heading("3.4 Clinical NLP", 2)
add_table(
    ["Feature", "Endpoint", "Description"],
    [
        ["Clinical Notes NER",     "/analyze-notes",     "NER for drugs, diagnoses, symptoms (spaCy + scispaCy)"],
        ["ICD-10 Auto-coding",     "/analyze-notes",     "Rule-based ICD-10 code suggestion"],
        ["Drug Interaction Check", "/drug-interactions", "12-pair knowledge base (major/moderate severity)"],
        ["Medication Extraction",  "/ai/medication-plan","Extract dose/frequency from free-text prescriptions"],
        ["Patient Trajectory",     "/analyze-notes",     "Classify note as improving/stable/deteriorating"],
    ],
    [1.8, 1.8, 3.0]
)

heading("3.5 Vital Signs Forecasting", 2)
bullet("LSTM neural network for sequence prediction")
bullet("Facebook Prophet for trend decomposition")
bullet("MEWS (Modified Early Warning Score): low / moderate / high / critical")
bullet("Supported vitals: heart rate, systolic BP, respiratory rate, temperature, SpO2")

heading("3.6 ARIA Ambulance Dispatch", 2)
bullet("Live Leaflet map: ambulance positions, status (available/en_route/at_scene)")
bullet("Claude-powered dispatch: selects nearest ambulance + best hospital by bed count and specialty")
bullet("Driver map with patient alert overlays and route polyline")
bullet("Dispatch history with resolution/cancellation tracking")

heading("3.7 MCI Triage Board", 2)
bullet("Voice-to-triage: Claude extracts vitals + category (P1/P2/P3/P4) from free text")
bullet("Stateless triage scoring: instant score from vitals without DB write")
bullet("Hospital allocation: suggests best receiving hospital given triage category")

heading("3.8 Medication Reminders", 2)
para("Per-patient medication schedule. Paste prescription text; Claude extracts drug, dose, frequency, timing. Reminders stored in MongoDB.")
doc.add_page_break()

# ── 4. API REFERENCE ──────────────────────────────────────────────────────────
heading("4. API Reference", 1, (79, 70, 229))
heading("4.1 Backend Routes (port 5000)", 2)
add_table(
    ["Method", "Path", "Description"],
    [
        ["POST", "/api/data/upload",           "Upload patient CSV dataset"],
        ["GET",  "/api/data",                  "List stored datasets"],
        ["POST", "/api/cluster",               "Run clustering algorithm"],
        ["GET",  "/api/cluster/results",       "List all cluster results"],
        ["POST", "/api/cluster/predict",       "Predict risk tier for new patient"],
        ["GET",  "/api/media/models",          "List imaging models (incl. LungLens)"],
        ["POST", "/api/media/upload",          "Upload medical image / DICOM"],
        ["POST", "/api/media/analyze/:fileId", "Run imaging analysis"],
        ["POST", "/api/media/explain/:fileId", "Claude AI explanation of findings"],
        ["POST", "/api/media/ai-chat",         "CliniQ multilingual vision chat"],
        ["POST", "/api/ai/cluster-insights",   "Claude cohort insights"],
        ["POST", "/api/ai/medication-plan",    "Extract medication schedule from text"],
        ["POST", "/api/dispatch/emergency",    "Submit emergency, get ARIA decision"],
        ["POST", "/api/triage/voice",          "Voice/text to triage via Claude"],
        ["GET",  "/api/health",                "Health check"],
    ],
    [0.7, 2.5, 3.4]
)
heading("4.2 ML Engine Routes (port 8000)", 2)
add_table(
    ["Method", "Path", "Description"],
    [
        ["POST", "/cluster",           "Run clustering on preprocessed data"],
        ["POST", "/optimal-k",         "Auto-select best K (Elbow + Silhouette)"],
        ["POST", "/feature-importance","Rank features by MI + ANOVA F-test"],
        ["POST", "/reduce-dimensions", "UMAP / t-SNE 2D projection"],
        ["POST", "/detect-anomalies",  "Isolation Forest anomaly detection"],
        ["POST", "/explain",           "SHAP feature importance per patient"],
        ["POST", "/analyze-notes",     "Clinical notes NLP pipeline"],
        ["POST", "/drug-interactions", "Drug interaction safety check"],
        ["POST", "/forecast-vitals",   "LSTM/Prophet vital forecast + MEWS"],
        ["GET",  "/models",            "List imaging models"],
        ["POST", "/analyze-image",     "DenseNet chest X-ray analysis"],
        ["POST", "/gradcam",           "Grad-CAM heatmap generation"],
        ["POST", "/ask",               "RAG clinical Q&A chatbot"],
    ],
    [0.7, 2.0, 3.9]
)
doc.add_page_break()

# ── 5. TECH STACK ─────────────────────────────────────────────────────────────
heading("5. Technology Stack", 1, (79, 70, 229))
add_table(
    ["Category", "Technology", "Version"],
    [
        ["Frontend Framework",  "React",            "18.2"],
        ["Build Tool",          "Vite",             "5.1"],
        ["Styling",             "Tailwind CSS",     "3.4"],
        ["Charts",              "Recharts + D3.js", "2.12 / 7.9"],
        ["HTTP Client",         "Axios",            "1.6"],
        ["Backend Runtime",     "Node.js",          "25.2"],
        ["Backend Framework",   "Express",          "4.18"],
        ["ODM",                 "Mongoose",         "8.1"],
        ["AI SDK",              "Anthropic SDK",    "0.95"],
        ["AI Model",            "Claude Sonnet 4.6","Anthropic"],
        ["ML Framework",        "scikit-learn",     "1.4"],
        ["Deep Learning",       "PyTorch",          "2.2"],
        ["Medical Imaging",     "torchxrayvision",  "1.0"],
        ["NLP",                 "spaCy + scispaCy", "3.7 / 0.5"],
        ["Transformers",        "HuggingFace",      "4.40"],
        ["Forecasting",         "Prophet",          "1.1"],
        ["Vector Store",        "FAISS",            "1.8"],
        ["Explainability",      "SHAP",             "0.45"],
        ["Dimensionality",      "UMAP-learn",       "0.5"],
        ["Database",            "MongoDB",          "7"],
    ],
    [2.0, 2.0, 2.6]
)
doc.add_page_break()

# ── 6. RUNNING THE PROJECT ────────────────────────────────────────────────────
heading("6. Running the Project", 1, (79, 70, 229))
para("Prerequisites: Node.js 18+, Python 3.11+, MongoDB running as a service.")

heading("6.1 Environment Variables (medicluster/backend/.env)", 2)
ep = doc.add_paragraph()
ep.paragraph_format.left_indent = Inches(0.5)
er = ep.add_run(
    "MONGO_URI=mongodb://localhost:27017/medicluster\n"
    "ML_ENGINE_URL=http://localhost:8000\n"
    "PORT=5000\n"
    "ANTHROPIC_API_KEY=<your_key>"
)
er.font.name = "Courier New"
er.font.size = Pt(9)

heading("6.2 Start All Services", 2)
add_table(
    ["Service", "Directory", "Command"],
    [
        ["ML Engine", "medicluster/ml-engine/", "python app.py"],
        ["Backend",   "medicluster/backend/",   "npm start"],
        ["Frontend",  "medicluster/frontend/",  "npm run dev"],
    ],
    [1.2, 2.5, 2.9]
)
para("Open http://localhost:3000 in your browser.")

# ── 7. RECENT ADDITIONS ───────────────────────────────────────────────────────
heading("7. Recent Feature Additions", 1, (79, 70, 229))
add_table(
    ["Feature", "Description", "Files Changed"],
    [
        ["Multilingual CliniQ",
         "Language selector (EN/HI/KN/TE/TA). Claude responds entirely in chosen language.",
         "AskAIPage.jsx, apiClient.js, mediaRoutes.js"],
        ["LungLens Model",
         "Claude Vision model on Imaging page. Returns same structured findings as DenseNet.",
         "mediaRoutes.js"],
        ["Shared Anthropic Client",
         "Single SDK client with 120s timeout. Fixes APIConnectionTimeoutError on large images.",
         "mediaRoutes.js"],
        ["Ambulance Dispatch Maps",
         "Live Leaflet maps for dispatcher and driver with route polylines and patient alerts.",
         "DispatchMap.jsx, DriverMap.jsx"],
    ],
    [1.6, 3.0, 2.0]
)

out = r"c:\Users\Admin\Desktop\Patient Health Risk Segregation\MediCluster_Documentation.docx"
doc.save(out)
print("Saved:", out)
